from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]


def set_normal_style(doc: Document, font_name: str = "Arial", font_size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    # python-docx uses built-in headings
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def add_code_block(doc: Document, code: str) -> None:
    # Simple code formatting: a single paragraph with monospaced font
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def build_main_theory_doc() -> Document:
    doc = Document()
    set_normal_style(doc)
    add_title(doc, "Теория по проекту: Full-Stack Quiz Application (LR0–LR7)")

    add_section_heading(doc, "1. Общая идея проекта", level=1)
    add_paragraph(
        doc,
        "К концу курса создаётся полноценное full-stack приложение Quiz: фронтенд на React + TypeScript, "
        "который общается с API бэкенда, а также поддержка offline-first (PWA) и тестирование. "
        "По структуре приложение делится на три слоя: фронтенд, backend API и база данных.",
    )
    add_bullets(
        doc,
        [
            "Frontend (LR1–7, 12): React + TypeScript + Vite + Tailwind CSS.",
            "State management: MobX (бизнес-логика) и Zustand (UI-состояние).",
            "Server state и синхронизация: React Query.",
            "Backend (в последующих лабораторных): Node.js + Hono + Prisma + Zod + JWT.",
            "Контракт между слоями: OpenAPI/Swagger и кодогенерация клиента.",
            "Качество: unit/component/E2E тесты (Vitest, React Testing Library, Playwright).",
            "Offline-first: Service Workers и стратегии кэширования.",
        ],
    )

    add_section_heading(doc, "2. Frontend-слой (что и зачем)", level=1)
    add_section_heading(doc, "2.1 TypeScript как основа типобезопасности", level=2)
    add_paragraph(doc, "TypeScript добавляет систему типов поверх JavaScript: компилятор помогает ловить ошибки до запуска и даёт автодополнение.")
    add_bullets(
        doc,
        [
            "Ключевая идея: «описываем форму данных и поведение» — и TypeScript снижает риск runtime-ошибок.",
            "Составные типы: union/intersection — для безопасной работы с вариантами данных.",
            "Интерфейсы и `type` — для описания моделей и контрактов с API.",
            "Generics — для переиспользуемых типобезопасных утилит и компонентов.",
            "DOM API и event handlers тоже типизируются: уменьшается число ошибок с неправильными типами элементов/событий.",
        ],
    )

    add_section_heading(doc, "2.2 React + TypeScript: типизация компонентов и хуков", level=2)
    add_bullets(
        doc,
        [
            "Функциональные компоненты типизируются через `props` и корректные типы для `children` и callback-props.",
            "Состояние (`useState`, `useReducer`) и эффекты (`useEffect`) имеют явные типы, особенно в сложных формах.",
            "Типизация обработчиков событий (event handlers) помогает избежать ошибок с `target/currentTarget`.",
            "Формы и валидация — часть корректного UX и надёжности (часто используют схемы валидации, например Zod).",
        ],
    )

    add_section_heading(doc, "2.3 Vite и Tailwind CSS", level=2)
    add_bullets(
        doc,
        [
            "Vite — быстрый dev-server и современная сборка, удобная разработка благодаря HMR.",
            "Tailwind — utility-first CSS: стили описываются атомарными классами, что ускоряет итерации и уменьшает «зоопарк» кастомных CSS правил.",
            "Совместная работа Vite + Tailwind обычно даёт быстрый цикл: правим UI → видим результат почти сразу.",
        ],
    )

    add_section_heading(doc, "2.4 Управление состоянием: MobX + Zustand", level=2)
    add_paragraph(doc, "В проекте удобно разделять ответственность: MobX — бизнес-логика и данные игры, Zustand — UI-настройки (темы, модалки, флаги).")
    add_section_heading(doc, "MobX (бизнес-логика)", level=3)
    add_bullets(
        doc,
        [
            "`observable` — наблюдаемое состояние (данные меняются → зависимые компоненты обновляются).",
            "`action` — «официальные» точки изменения состояния (особенно важно в строгих режимах).",
            "`computed` — вычисляемые значения (прогресс, статистика) с кэшированием пересчёта.",
            "В React используется `observer` (или `useObserver`) — компонент подписывается только на реально читаемые поля.",
        ],
    )
    add_section_heading(doc, "Zustand (UI-состояние)", level=3)
    add_bullets(
        doc,
        [
            "Простой API: один хук, созданный через `create`.",
            "Селекторы: подписка компонента только на нужную часть store уменьшает лишние перерисовки.",
            "Middleware: `persist` для сохранения настроек (например темы) в `localStorage`.",
            "Async actions можно писать привычно через async/await (в отличие от реактивной модели MobX).",
        ],
    )

    add_section_heading(doc, "2.5 React Query: server state без boilerplate", level=2)
    add_paragraph(doc, "React Query решает проблему «ручного fetch»: кэширование, дедупликация запросов, обработка loading/error и удобные мутации.")
    add_bullets(
        doc,
        [
            "`useQuery` — для чтения данных (GET).",
            "`useMutation` — для изменений (POST/PUT/DELETE).",
            "`queryKey` — ключи кэша: по ним React Query понимает, что и когда пере-запрашивать.",
            "Mutations дают состояния (`isPending`, `isError`, `isSuccess`) — на них строятся кнопки и UX.",
            "В связке с OpenAPI/кодогенерацией клиент становится типобезопасным и быстро интегрируется.",
        ],
    )

    add_section_heading(doc, "2.6 React-паттерны и надёжность UI", level=2)
    add_paragraph(doc, "Хороший код — это не только «чтобы работало», но и чтобы приложение было устойчивым и поддерживаемым.")
    add_bullets(
        doc,
        [
            "Custom Hooks — вынесение stateful-логики в переиспользуемые функции.",
            "Compound Components — композиция сложного UI через части (Header/Body/Footer и т.п.).",
            "Context API — избегание prop drilling для действительно «глобальных» данных (theme/auth).",
            "Error Boundaries — ловят ошибки при рендере компонентов и предотвращают «white screen of death».",
            "Важно: Error Boundary не ловит ошибки в обработчиках событий и асинхронных запросах — их обрабатывают отдельно.",
        ],
    )

    add_section_heading(doc, "3. API-уровень: HTTP, REST, валидация и безопасность", level=1)
    add_paragraph(doc, "В проекте фронтенд и бэкенд общаются через REST API. Вся архитектура опирается на HTTP-методы, статусы, заголовки и контракт данных.")
    add_section_heading(doc, "3.1 REST и CRUD→HTTP", level=2)
    add_bullets(
        doc,
        [
            "Create → `POST /resources`.",
            "Read → `GET /resources/:id`.",
            "Update → `PUT` или `PATCH` (семантика частичного изменения).",
            "Delete → `DELETE /resources/:id`.",
            "REST подразумевает stateless коммуникацию: клиент передаёт нужные данные в запросе.",
        ],
    )
    add_section_heading(doc, "3.2 Коды ответов и заголовки", level=2)
    add_bullets(
        doc,
        [
            "2xx — успех (200, 201, 204).",
            "3xx — редиректы (301, 302, 304).",
            "4xx — ошибки клиента (400, 401, 403, 404, 422, 429).",
            "5xx — ошибки сервера (500, 502, 503).",
            "Headers: `Content-Type` (формат тела), `Authorization` (Bearer/JWT), `Cache-Control`, `CORS` (политика доступа).",
        ],
    )
    add_section_heading(doc, "3.3 Валидация данных и формат ошибок", level=2)
    add_bullets(
        doc,
        [
            "Нужна и клиентская, и серверная валидация: UX на клиенте + безопасность на сервере.",
            "Для схем/контрактов часто используют runtime-валидацию (например Zod) и возвращают понятные ошибки (в т.ч. для `422`).",
            "Типы в TypeScript дают статическую уверенность, но runtime проверка всё равно нужна, потому что данные приходят снаружи.",
        ],
    )
    add_section_heading(doc, "3.4 AAA: Authentication / Authorization / Accounting", level=2)
    add_bullets(
        doc,
        [
            "Authentication: «кто ты?» (JWT/cookies/session).",
            "Authorization: «что тебе можно?» (роль/права: student/admin).",
            "Accounting: учёт действий (обычно логирование/аудит).",
            "В проекте используется Bearer token (JWT): токен передаётся в заголовке `Authorization: Bearer <token>`.",
        ],
    )

    add_section_heading(doc, "4. OpenAPI и кодогенерация клиента", level=1)
    add_paragraph(
        doc,
        "OpenAPI описывает контракт REST API (эндпоинты, схемы запросов/ответов). "
        "Кодогенерация (например через `orval`) превращает спецификацию в типобезопасные хуки React Query.",
    )
    add_bullets(
        doc,
        [
            "Проблема ручной типизации: backend может поменяться, а frontend продолжит компилироваться (и падать в runtime).",
            "OpenAPI помогает держать фронт и бэк синхронизированными.",
            "Кодогенерация даёт: типы, хуки (`useQuery`/`useMutation`) и единый стиль работы с API.",
            "Workflow обычно такой: обновили OpenAPI → запустили `npm run codegen` → использовали хуки в компонентах.",
        ],
    )

    add_section_heading(doc, "5. Тестирование (Vitest, RTL, Playwright)", level=1)
    add_bullets(
        doc,
        [
            "Тестовая пирамида: unit (база), integration (связи), e2e (минимум, но максимально ценно).",
            "Vitest подходит для быстрых unit/component тестов в экосистеме Vite.",
            "React Testing Library учит тестировать «как пользователь» (queries по роли/лейблу/тексту).",
            "Playwright делает end-to-end сценарии: полный пользовательский flow через браузер.",
            "Coverage — не единственный показатель качества, но порог в проекте обычно держат на уровне ~70%.",
        ],
    )

    add_section_heading(doc, "6. Методология лабораторных работ", level=1)
    add_bullets(
        doc,
        [
            "Fork репозитория курса → работа в своей ветке.",
            "Решение задания (в классе или дома).",
            "Оформление и отправка Pull Request для проверки.",
            "Подготовка: использовать материалы `lr*/docs/`, интерактивные руководства и примеры кода.",
        ],
    )

    add_section_heading(doc, "7. Короткая привязка теории к LR0–LR7", level=1)
    add_bullets(
        doc,
        [
            "LR0: базовые навыки HTML/CSS/JavaScript (практика DOM и вёрстки).",
            "LR1: TypeScript (типы, union/intersection, интерфейсы, generics, tsconfig).",
            "LR2: React + TypeScript (типизация props/state/hooks, формы, события, паттерны).",
            "LR4: state management (MobX и Zustand + разделение бизнес/UI логики).",
            "LR5/LR6or5: интеграция с API через React Query и типобезопасные хуки из OpenAPI.",
            "LR7: интеграция Quiz с API (создание сессии, отправка ответов, завершение сессии).",
        ],
    )

    return doc


def build_lr7_execution_doc() -> Document:
    doc = Document()
    set_normal_style(doc)
    add_title(doc, "LR7: Теория по выполнению лабораторной работы (интеграция Quiz с API)")

    add_section_heading(doc, "1. Суть задания LR7", level=1)
    add_paragraph(
        doc,
        "В LR7 вы превращаете локальную логику квиза в сценарий, управляемый backend API: "
        "создаёте игровую сессию, отправляете ответы и завершаете сессию, получая результаты.",
    )
    add_bullets(
        doc,
        [
            "Фронтенд использует React Query мутации для POST-запросов.",
            "Хуки API генерируются из OpenAPI схемы (через `orval`).",
            "Для выполнения достаточно 3 API-хука: создание сессии, отправка ответа, завершение сессии.",
        ],
    )

    add_section_heading(doc, "2. Что у вас уже подготовлено", level=1)
    add_bullets(
        doc,
        [
            "Скелет приложения: кнопка `Login/Logout` и базовая авторизация.",
            "Настроенная генерация API-хуков из `quiz-api-schema.yaml` (см. `orval.config.ts`).",
            "Mock-сервер для разработки (запускается на `http://localhost:3000`).",
        ],
    )

    add_section_heading(doc, "3. Подготовка окружения", level=1)
    add_paragraph(doc, "Чтобы синхронно проверять фронтенд и API, нужно запустить оба сервиса: mock backend и dev-сервер приложения.")
    add_section_heading(doc, "3.1 Запуск mock-сервера", level=2)
    add_code_block(doc, "cd mock-server")
    add_code_block(doc, "npm install")
    add_code_block(doc, "npm start")
    add_paragraph(doc, "Mock API доступен по адресу `http://localhost:3000`. Токен для авторизации можно получить через mock-эндпоинт OAuth callback.")

    add_section_heading(doc, "3.2 Запуск приложения", level=2)
    add_code_block(doc, "npm run dev")
    add_paragraph(doc, "Обычно приложение открывается по адресу `http://localhost:5173`.")

    add_section_heading(doc, "4. Перенос реализации квиза из LR4", level=1)
    add_paragraph(
        doc,
        "LR7 предполагает, что базовый UI квиза уже реализован. Поэтому нужно перенести готовые файлы (Task4 и связанные stores/types/questions) из LR4."
    )
    add_bullets(
        doc,
        [
            "Task4.tsx (компонент квиза).",
            "gameStore.ts и uiStore.ts (состояние игры и интерфейсные флаги).",
            "types/quiz.ts (модели данных).",
            "data/questions.ts (локальные вопросы/заготовки).",
        ],
    )

    add_section_heading(doc, "5. Множественный выбор: что именно нужно поменять", level=1)
    add_paragraph(doc, "В LR7 важный шаг — поддержка multiple-select: пользователь может выбрать несколько вариантов ответа, а сервер проверяет их и начисляет баллы/штрафы.")

    add_section_heading(doc, "5.1 Измените store", level=2)
    add_bullets(
        doc,
        [
            "Замените `selectedAnswer: number | null` на массив `selectedAnswers: number[]` (или аналогичный тип).",
            "Добавьте метод `toggleAnswer(answerIndex: number)`, который включает выбранный вариант или убирает его при повторном клике.",
            "Обновите все места, где использовался `selectedAnswer`, чтобы они читали `selectedAnswers`.",
        ],
    )

    add_section_heading(doc, "5.2 Измените компонент Task4.tsx", level=2)
    add_bullets(
        doc,
        [
            "Клик по варианту должен вызывать `gameStore.toggleAnswer(index)`.",
            "Логика кнопки «Далее/Следующий» должна ориентироваться на количество выбранных ответов.",
            "В UI можно показывать галочку `✓` для выбранного варианта вместо буквы `A/B/C...`.",
        ],
    )

    add_section_heading(doc, "6. Кодогенерация API-хуков (OpenAPI -> React Query)", level=1)
    add_paragraph(doc, "API не нужно описывать вручную в сотнях мест. Кодогенерация создаёт типобезопасные хуки на основе OpenAPI спецификации.")
    add_bullets(
        doc,
        [
            "Запустите `npm run codegen` (внутри обычно вызывается `orval`).",
            "Проверьте, что сгенерировались хуки в директории `generated/api`.",
            "После генерации обратите внимание на override в `orval.config.ts`: используется кастомная функция `customFetch` из `src/api/client.ts`.",
        ],
    )
    add_paragraph(doc, "Если OpenAPI обновлялся, кодогенерацию обычно повторяют перед тем как компилировать/тестировать.")

    add_section_heading(doc, "7. Интеграция с API: сценарий мутаций", level=1)
    add_paragraph(doc, "В LR7 вы управляете квизом через последовательность запросов к API. Основная идея: state store содержит текущий UI-статус, а сервер — авторитет по проверке и расчётам.")

    add_section_heading(doc, "7.1 Создание сессии", level=2)
    add_bullets(
        doc,
        [
            "Используйте хук `usePostApiSessions` для создания игровой сессии.",
            "При старте игры вызовите mutate с `questionCount` и `difficulty` (а в реальном сценарии — ещё и фильтрами/категориями).",
            "В `onSuccess` сохраните `sessionId` и загрузите вопросы в `gameStore` (через метод вроде `setQuestionsFromAPI`).",
        ],
    )
    add_section_heading(doc, "7.2 Отправка ответа на вопрос", level=2)
    add_bullets(
        doc,
        [
            "Используйте `usePostApiSessionsSessionIdAnswers` при переходе к следующему вопросу.",
            "Отправляйте на сервер `questionId` и данные выбранного ответа:",
            "для multiple-select — массив выбранных индексов (selected options).",
            "после ответа обновляйте результат в store (pointsEarned/status/feedback) и только потом переходите к следующему вопросу.",
            "Ключевой UX: кнопку «Далее» стоит блокировать, пока мутация в `isPending`, чтобы не отправлять дубликаты.",
        ],
    )

    add_section_heading(doc, "7.3 Завершение сессии", level=2)
    add_bullets(
        doc,
        [
            "Используйте `usePostApiSessionsSessionIdSubmit` на последнем шаге («Завершить»).",
            "После завершения обновите UI: отметьте игру как законченной и покажите результаты.",
            "Если sessionId отсутствует (нестандартные сценарии), fallback — завершить локально.",
        ],
    )

    add_section_heading(doc, "8. Почему типы `id` часто конфликтуют", level=1)
    add_paragraph(
        doc,
        "В разных источниках (mock/db/schema) идентификаторы могут быть строковыми или числовыми. "
        "Если TypeScript ожидает `number`, а API возвращает `string`, вы получите ошибку компиляции. "
        "Поэтому в `src/types/quiz.ts` в LR7 обычно допускают `string | number`."
    )
    add_code_block(doc, "export interface Question { id: string | number; ... }")
    add_code_block(doc, "export interface Answer { questionId: string | number; ... }")

    add_section_heading(doc, "9. Проверка результата", level=1)
    add_bullets(
        doc,
        [
            "Login: токен появляется/обновляется в `localStorage` (или используется как Bearer для запросов).",
            "Старт игры: загружаются вопросы с сервера и появляются в UI.",
            "Выбор ответов: выбранные варианты визуально отмечаются.",
            "Следующий вопрос: запрос `POST /api/sessions/{id}/answers` реально уходит в Network.",
            "Завершение: на последнем вопросе выполняется `POST /api/sessions/{id}/submit` и появляется итоговая статистика.",
            "В DevTools: отсутствуют ошибки в Console, а в React Query DevTools видны мутации и кэшированные данные.",
        ],
    )

    add_section_heading(doc, "10. Типовые ошибки и что делать", level=1)
    add_bullets(
        doc,
        [
            "Ошибка вида `Type 'string' is not assignable to type 'number'`: расширьте типы `id/questionId` до `string | number` или приведите данные к одному формату.",
            "Ошибки из-за `null` (например, `Cannot read property 'id' of null`): добавляйте защиту перед использованием данных (проверки на `currentQuestion`).",
            "Повторные запросы при быстрой кликабельности: проверяйте `disabled={isPending}` и корректно сбрасывайте/фиксируйте состояние выбранных ответов.",
        ],
    )

    add_section_heading(doc, "11. Теоретическое «как это работает» (упрощённо)", level=1)
    add_paragraph(
        doc,
        "Суммарно процесс выглядит так: UI фиксирует выбор пользователя, затем отправляет его на сервер. "
        "Сервер проверяет ответ (для multiple-select — автоматически по penalty/минимальному порогу и формирует `pointsEarned/status/breakdown`; "
        "для essay — сохраняет как ожидающий проверки). "
        "После этого фронтенд обновляет состояние игры и ведёт пользователя к следующему вопросу."
    )

    return doc


def main() -> None:
    out_main = ROOT / "Теория_проекта_LR0-LR7.docx"
    out_lr7 = ROOT / "Теория_по_выполнению_LR7.docx"

    doc_main = build_main_theory_doc()
    doc_main.save(out_main)

    doc_lr7 = build_lr7_execution_doc()
    doc_lr7.save(out_lr7)

    print(f"Saved: {out_main}")
    print(f"Saved: {out_lr7}")


if __name__ == "__main__":
    main()

